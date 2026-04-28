import re
import os
from typing import Optional
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    title: str
    content: str
    summary: str


class ParserService:
    """文件解析服务，支持 PDF 和 Word 文档"""

    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        """
        根据文件类型调用对应的解析器

        Args:
            file_path: 文件路径
            file_type: 文件类型 (pdf/docx/doc)

        Returns:
            ParsedDocument: 解析后的文档对象
        """
        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return self._parse_word(file_path, file_type)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件"""
        try:
            import pdfplumber

            text_parts = []
            title = ""

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    if i == 0:
                        # 从第一页提取标题（取前100字符或第一行）
                        lines = page_text.split("\n")
                        title = lines[0][:200].strip() if lines else ""
                    text_parts.append(page_text)

            content = "\n".join(text_parts)
            summary = self._generate_summary(content)

            if not title:
                title = self._extract_title_from_content(content) or "未命名文档"

            return ParsedDocument(
                title=title,
                content=content,
                summary=summary
            )
        except Exception as e:
            raise Exception(f"PDF 解析失败: {str(e)}")

    def _parse_word(self, file_path: str, file_type: str) -> ParsedDocument:
        """解析 Word 文件"""
        # 从文件路径提取原始文件名作为备用标题
        import os
        original_filename = os.path.splitext(os.path.basename(file_path))[0]

        # 先检查是否有 chunk.xhtml（PDF转Word残留的正文，内容更完整）
        xhtml_content = self._parse_docx_chunk(file_path)
        if xhtml_content and len(xhtml_content) > 200:  # chunk.xhtml通常有较多内容
            title = self._extract_title_from_content(xhtml_content) or original_filename or "未命名文档"
            summary = self._generate_summary(xhtml_content)
            return ParsedDocument(
                title=title,
                content=xhtml_content,
                summary=summary
            )

        # 再尝试用 python-docx 解析
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            # 过滤掉全是点的无效内容
            valid_content = self._filter_valid_content(content)

            if valid_content.strip():
                title = paragraphs[0][:200] if paragraphs else "未命名文档"
                summary = self._generate_summary(valid_content)
                return ParsedDocument(
                    title=title,
                    content=valid_content,
                    summary=summary
                )
        except Exception as e:
            pass  # 继续尝试其他方法

        # .doc 格式用 python-docx 解析可能失败，尝试用 Windows COM 自动化
        if file_type == "doc":
            try:
                content = self._parse_doc_with_word_com(file_path)
                if content.strip():
                    lines = content.split("\n")
                    title = lines[0][:200].strip() if lines else "未命名文档"
                    summary = self._generate_summary(content)
                    return ParsedDocument(
                        title=title,
                        content=content,
                        summary=summary
                    )
            except Exception as e:
                raise Exception(f"Word 解析失败: {str(e)}")

        raise Exception("Word 文档解析失败，请转换为 .docx 格式后重试")

    def _filter_valid_content(self, content: str) -> str:
        """过滤掉全是点或无效符号的内容"""
        if not content:
            return ""
        # 过滤掉全是由"."组成的行
        lines = content.split("\n")
        valid_lines = []
        for line in lines:
            stripped = line.strip()
            # 如果一行全是点（超过10个），跳过
            if stripped and not (len(stripped) > 10 and all(c == '.' for c in stripped)):
                valid_lines.append(line)
        return "\n".join(valid_lines)

    def _parse_docx_chunk(self, file_path: str) -> str:
        """从 docx 中提取 chunk.xhtml 的内容（PDF转Word残留）"""
        try:
            import zipfile
            from html.parser import HTMLParser

            with zipfile.ZipFile(file_path) as z:
                if 'chunk.xhtml' in z.namelist():
                    raw = z.read('chunk.xhtml')
                    # 尝试多种编码，UTF-8优先，失败则尝试GBK
                    xhtml = None
                    for encoding in ('utf-8', 'gbk', 'gb2312', 'latin-1'):
                        try:
                            xhtml = raw.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    if xhtml is None:
                        xhtml = raw.decode('utf-8', errors='replace')
                    # 解析HTML提取纯文本
                    return self._strip_html_tags(xhtml)
        except Exception:
            pass
        return ""

    def _strip_html_tags(self, html: str) -> str:
        """去掉HTML标签提取纯文本"""
        from html.parser import HTMLParser

        class TagStripper(HTMLParser):
            def __init__(self):
                super().__init__()
                self.result = []
                self.skip_data = False

            def handle_starttag(self, tag, attrs):
                # 跳过script和style
                if tag in ('script', 'style'):
                    self.skip_data = True

            def handle_endtag(self, tag):
                if tag in ('script', 'style'):
                    self.skip_data = False

            def handle_data(self, data):
                if not self.skip_data:
                    self.result.append(data)

            def get_text(self):
                text = ''.join(self.result)
                # 清理多余空白
                import re
                text = re.sub(r'\s+', ' ', text)
                return text.strip()

        stripper = TagStripper()
        stripper.feed(html)
        return stripper.get_text()

    def _parse_doc_with_word_com(self, file_path: str) -> str:
        """使用 Windows COM 自动化调用 Word 解析 .doc 文件"""
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(os.path.abspath(file_path))
                content = doc.Content.Text
                doc.Close(False)
                return content
            finally:
                word.Quit()
        finally:
            pythoncom.CoUninitialize()

    def _generate_summary(self, content: str, max_length: int = 500) -> str:
        """生成摘要，取内容前 max_length 个字符"""
        content = content.replace("\n", " ").replace("\r", " ")
        content = re.sub(r"\s+", " ", content)
        if len(content) <= max_length:
            return content
        return content[:max_length] + "..."

    def _extract_title_from_content(self, content: str) -> Optional[str]:
        """从内容中提取标题"""
        lines = content.split("\n")
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 5 and len(line) < 200:
                return line
        return None


parser_service = ParserService()
